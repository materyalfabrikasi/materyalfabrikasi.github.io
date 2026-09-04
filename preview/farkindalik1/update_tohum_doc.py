from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE = r"C:\Users\Tugba Gulsen HUR\Desktop\TOHUM.docx"
OUTPUT = r"C:\Users\Tugba Gulsen HUR\Desktop\TOHUM - Leylo Etkinliği.docx"
LEYLO = r"C:\Users\Tugba Gulsen HUR\Desktop\leylo\leylo.png"
REMEMBER = r"C:\Users\Tugba Gulsen HUR\Desktop\leylo\hatirliyor-musun.png"

doc = Document(SOURCE)

def style_run(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_text(paragraph, text, size=11, bold=False, color=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph

def add_before(anchor, text, style=None, size=11, bold=False, color=None):
    paragraph = anchor.insert_paragraph_before(style=style)
    set_text(paragraph, text, size=size, bold=bold, color=color)
    return paragraph

def apply_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)

# Genel bilgi alanlarını çalışmaya göre tamamla.
set_text(doc.paragraphs[10], "İlişkilendirilen Tema: Tanışma, selamlaşma ve kendini ifade etme", 11)
set_text(doc.paragraphs[11], "İlgili Hafta: Okul öncesi tanışma haftası", 11)

paragraphs = doc.paragraphs
set_text(paragraphs[13], "ETKİNLİK: LEYLO İLE MERHABA VE BENİM HAREKETİM", 15, True, (52, 71, 43))
set_text(paragraphs[14], "Yaş grubu: 5-6 yaş  |  Süre: 20-25 dakika  |  Uygulama: Bireysel ve grup katılımı", 10, True)
set_text(paragraphs[15], "Amaç: Öğrencilerin “merhaba” ifadesini farklı ses ve karakterlerle doğal biçimde duymaları; adlarını söylemeleri, bir hareket seçmeleri ve arkadaşlarının adlarıyla hareketlerini hatırlamalarıdır.", 11)

# Şablondaki boş alana iki temel oyun görselini yerleştir.
visual = paragraphs[16]
visual.clear()
visual.alignment = WD_ALIGN_PARAGRAPH.CENTER
visual.paragraph_format.space_after = Pt(8)
visual.add_run().add_picture(LEYLO, width=Inches(1.65))
visual.add_run("     ")
visual.add_run().add_picture(REMEMBER, width=Inches(2.05))

teacher_note = next(p for p in doc.paragraphs if "NOT:" in p.text.upper())

for item in [
    "Selamlaşma ifadesini dinler ve uygun bağlamda tekrar eder.",
    "Kendi adını anlaşılır biçimde söylemeye istek gösterir.",
    "Görsel ve işitsel uyaranlara dikkatini yöneltir.",
    "Bir hareket seçer, uygular ve grup içinde tekrar eder.",
    "Arkadaşının adını ve yaptığı hareketi hatırlamaya çalışır.",
]:
    add_before(teacher_note, item, size=10)

add_before(teacher_note, "Dijital etkinlik akışı", size=12, bold=True, color=(185, 38, 31))
steps = [
    ("Karşılama", "Öğrenci açılış ekranındaki “Başlat” düğmesine dokunur ve karşılama videosunu izler."),
    ("Leylo ve arkadaşları", "Leylo ortada; robot, tavşan, kuş ve ayı çevresinde yer alır. Bir karaktere dokunulduğunda karakter büyür, kendi “merhaba” sesi çalar ve ardından normal boyutuna döner."),
    ("Farklı söyleyişleri keşfetme", "Çocuk, karakterlerin farklı hız ve tonlardaki selamlaşmalarını dinler. Her dinleme bir başarı deneyimi olarak kabul edilir; doğru-yanlış değerlendirmesi yapılmaz."),
    ("Adını söyle", "Öğrenci görseli ekrana gelir. Kerem örnek sesiyle “Merhaba, ben Kerem” ifadesi duyulur. Ardından öğrenci kendi adını söyler."),
    ("Hareketini seç ve yap", "Öğrenci el sallama, alkışlama, kollarını açma veya zıplama görsellerinden birine dokunur. Seçilen görsel büyür ve harekete özgü bir efekt oynar. Öğrenci adını söyleyip hareketini yapar; sınıf onun adını söyleyerek aynı hareketi tekrar eder."),
    ("Hatırlıyor musun?", "Öğretmen bir öğrencinin hareketini tekrarlar. Sınıf hareketin kime ait olduğunu tahmin eder ve arkadaşının adını söyler."),
]
for title, description in steps:
    p = add_before(teacher_note, "", size=10)
    p.clear()
    r1 = p.add_run(f"{title}: ")
    style_run(r1, size=10, bold=True, color=(52, 71, 43))
    r2 = p.add_run(description)
    style_run(r2, size=10)
    p.paragraph_format.space_after = Pt(5)
    apply_numbering(p)

add_before(teacher_note, "Materyaller ve teknik hazırlık", size=12, bold=True, color=(185, 38, 31))
for label, item in [
    ("Araçlar", "Bilgisayar veya etkileşimli tahta, hoparlör ve güncel bir internet tarayıcısı."),
    ("Dosya", "leylo klasöründeki index.html dosyası tarayıcıda açılır; görsel, video ve ses dosyaları aynı klasörde tutulur."),
    ("Güvenlik", "Zıplama etkinliği öncesinde öğrencilerin çevresinde güvenli ve boş bir alan oluşturulur."),
]:
    p = add_before(teacher_note, "", size=10)
    r1 = p.add_run(f"{label}: ")
    style_run(r1, size=10, bold=True, color=(52, 71, 43))
    style_run(p.add_run(item), size=10)

# Taslak yer tutucu öğretmen notunu kullanıma hazır bir notla değiştir.
placeholder = next((p for p in doc.paragraphs if "PRATİK NOTLAR" in p.text.upper()), None)
if placeholder:
    set_text(placeholder, "Uygulama sırasında çocukların farklı hızlarda katılım gösterebileceği unutulmamalıdır. Söylemek veya hareket etmek istemeyen çocuk gözlemci olarak kalabilir ve hazır olduğunda etkinliğe katılabilir. Telaffuz doğruluğu puanlanmamalı; her katılım olumlu geri bildirimle desteklenmelidir.", 10, False, (52, 71, 43))

doc.save(OUTPUT)
print(OUTPUT)
